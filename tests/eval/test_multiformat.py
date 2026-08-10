"""Multi-format ingestion benchmark, end to end (task #30).

Proves every working format family ingests through its parser and its carried
fact is retrievable via the real hybrid pipeline. Uses the real embedder (the
benchmark is meaningless without dense retrieval), so it skips without the
`datum[embed]` extra. The formats not runnable in this environment (pdf,
image, audio) are asserted to be REPORTED as skipped-with-reason, so the
benchmark can never quietly claim coverage it does not have.
"""

from __future__ import annotations

import importlib.util
import os

import psycopg
import pytest

from datum.corpus import Corpus
from datum.eval.multiformat_benchmark import run_benchmark

_DSN = os.environ.get("DATUM_PG_DSN", "postgresql://localhost/datum")


def _pg_reachable(dsn: str) -> bool:
    try:
        with psycopg.connect(dsn, connect_timeout=2):
            return True
    except psycopg.OperationalError:
        return False


pytestmark = [
    pytest.mark.skipif(
        not _pg_reachable(_DSN), reason=f"no reachable Postgres at DATUM_PG_DSN={_DSN!r}"
    ),
    pytest.mark.skipif(
        importlib.util.find_spec("sentence_transformers") is None,
        reason="the datum[embed] extra is not installed; the multi-format benchmark needs the embedder",
    ),
    pytest.mark.skipif(
        importlib.util.find_spec("docling") is None,
        reason="the datum[parse] extra (docling) is not installed",
    ),
]


@pytest.fixture
def corpus():
    c = Corpus.open(_DSN, hit_signing_key=b"multiformat-key")
    with psycopg.connect(_DSN, autocommit=True) as conn:
        conn.execute("TRUNCATE TABLE records RESTART IDENTITY CASCADE")
        conn.execute("TRUNCATE TABLE wal_entries RESTART IDENTITY")
        conn.execute("TRUNCATE TABLE plan_traces")
        conn.execute("TRUNCATE TABLE view_cursors")
        conn.execute("DELETE FROM view_dense")
        conn.execute("DELETE FROM view_lexical")
    yield c
    c.close()


def test_every_working_format_ingests_and_retrieves(corpus, tmp_path):
    report = run_benchmark(corpus, tmp_path)

    # Every format that ran must have ingested AND retrieved its fact.
    failures = [f"{r.fmt}: {r.detail}" for r in report.results if not (r.ingested and r.retrieved)]
    assert report.passed, "multi-format benchmark regressed:\n" + "\n".join(failures)

    # The seven text-native families are actually covered (not silently empty).
    covered = {r.fmt for r in report.results}
    assert {"md", "txt", "html", "csv", "docx", "pptx", "xlsx"} <= covered

    # The blocked families are reported with a reason, never quietly dropped.
    skipped_fmts = {fmt for fmt, _ in report.skipped}
    assert {"pdf", "image/scanned", "audio"} <= skipped_fmts
    assert all(reason for _, reason in report.skipped)


def test_ingest_file_is_namespace_isolated_like_text_ingest(corpus, tmp_path):
    # A new ingestion entry point (ingest_file) must inherit the same tenancy
    # isolation as text ingest — a docling-ingested record in one namespace is
    # invisible to a principal in another (the compiler resolves the namespace
    # before any operator, and every operator re-checks it at the records join,
    # independent of HOW the record was written).
    import docx as _docx

    from datum.kernel.principal import Principal

    path = tmp_path / "secret.docx"
    d = _docx.Document()
    d.add_paragraph("The vault combination is seventeen forty-two ninety-nine.")
    d.save(path)
    corpus.ingest_file(str(path), Principal(id="a", namespace="tenant:acme"))

    # Give the outsider's OWN namespace matching content, so the abstention
    # floor is cleared by evil's own record and cannot mask the tenancy check:
    # if the records-join namespace filter were removed, evil's query WOULD
    # now surface acme's vault string. This is what makes the test
    # discriminate (verified by reverting ann_op's `r.namespace = %s`).
    evil_doc = tmp_path / "evil_vault.docx"
    d2 = _docx.Document()
    d2.add_paragraph("Our own vault combination here is three three three test decoy.")
    d2.save(evil_doc)
    corpus.ingest_file(str(evil_doc), Principal(id="b", namespace="tenant:evil"))

    ev_owner = corpus.search("what is the vault combination", principal=Principal(id="a", namespace="tenant:acme"))
    assert any("seventeen forty-two" in h.content for h in ev_owner.hits)

    ev_outsider = corpus.search("what is the vault combination", principal=Principal(id="b", namespace="tenant:evil"))
    assert ev_outsider.hits, "outsider should retrieve its OWN namespace's content (floor cleared)"
    assert all("seventeen forty-two" not in h.content for h in ev_outsider.hits), (
        "acme's vault string leaked into the outsider's results"
    )
